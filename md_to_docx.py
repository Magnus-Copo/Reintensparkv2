
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_file, docx_file):
    document = Document()
    
    # Set default style (optional, rely on defaults for now)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_content = []

    for line in lines:
        stripped_line = line.strip()
        
        # Code Block Handling
        if stripped_line.startswith('```'):
            if in_code_block:
                # End of code block
                p = document.add_paragraph()
                p.style = 'Quote' # Use Quote style for code background-ish feel or just normal keys
                runner = p.add_run('\n'.join(code_content))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(10)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line.rstrip())
            continue

        if not stripped_line:
            continue

        # Headers
        if stripped_line.startswith('# '):
            document.add_heading(stripped_line[2:], level=1)
        elif stripped_line.startswith('## '):
            document.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('### '):
            document.add_heading(stripped_line[4:], level=3)
        
        # List Items
        elif stripped_line.startswith('* ') or stripped_line.startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            format_text(p, stripped_line[2:])
        
        elif stripped_line.startswith('1. '):
            p = document.add_paragraph(style='List Number')
            format_text(p, stripped_line[3:])

        # Separator
        elif stripped_line == '---':
            document.add_paragraph('_' * 40) # Simple visual separator

        # Tables (Very basic handling)
        elif stripped_line.startswith('|'):
            # Ideally parse table, but for SOP.md, maybe just print as text or code?
            # Let's treat as simple paragraph for now or monospace
            p = document.add_paragraph()
            run = p.add_run(stripped_line)
            run.font.name = 'Courier New'

        # Normal Text
        else:
            p = document.add_paragraph()
            format_text(p, stripped_line)

    document.save(docx_file)

def format_text(paragraph, text):
    """
    Parses **bold**, *italic*, and `code` basics.
    This is a simple parser and doesn't handle nested or complex cases.
    """
    # Regex for bold: **text**
    # We need to split and handle parts.
    # Simple state machine or regex split?
    # Let's try splitting by **
    
    # This approach is simplistic: supports **bold** OR *italic* but not mix easily in one line properly with this logic
    # Improving: use regex to find tokens
    
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)
    
    for token in tokens:
        if not token:
            continue
            
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            # run.font.color.rgb = RGBColor(200, 50, 50) 
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            # Link: [text](url) - just show text for now, maybe bold it?
            match = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if match:
                link_text = match.group(1)
                # link_url = match.group(2)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                 paragraph.add_run(token)
        else:
            paragraph.add_run(token)

if __name__ == '__main__':
    md_path = 'SOP.md'
    docx_path = 'SOP.docx'
    print(f"Converting {md_path} to {docx_path}...")
    try:
        parse_markdown_to_docx(md_path, docx_path)
        print("Conversion successful.")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
